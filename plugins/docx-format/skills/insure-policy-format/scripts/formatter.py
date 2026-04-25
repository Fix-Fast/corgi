from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from _ooxml import patch_list_levels, set_list_suffix, strip_list_ind_overrides


LEVEL_PATTERNS = [
    re.compile(r"^\s*(\d+)[).](\s+|(?=[A-Z(]))"),
    re.compile(r"^\s*([a-z])[).](\s+|(?=[A-Z(]))"),
    re.compile(r"^\s*([ivx]+)[).](\s+|(?=[A-Z(]))"),
    re.compile(r"^\s*\((\d+)\)(\s+|(?=[A-Z(]))"),
    re.compile(r"^\s*\(([a-z])\)(\s+|(?=[A-Z(]))"),
    re.compile(r"^\s*\(([ivx]+)\)(\s+|(?=[A-Z(]))"),
]

EMBEDDED_MARKER_PATTERNS = [
    re.compile(r"(?<![A-Za-z0-9)])(\(\d+\)(\s+|(?=[A-Z(])))"),
    re.compile(r"(?<![A-Za-z0-9)])(\([a-z]\)(\s+|(?=[A-Z(])))"),
    re.compile(r"(?<![A-Za-z0-9)])(\([ivx]+\)(\s+|(?=[A-Z(])))"),
    re.compile(r"(?<=\s)(\d+\)(\s+|(?=[A-Z(])))"),
    re.compile(r"(?<=\s)([a-z]\)(\s+|(?=[A-Z(])))"),
    re.compile(r"(?<=\s)([ivx]+\)(\s+|(?=[A-Z(])))"),
]

LEVEL_TO_OL_ATTR = {
    0: [1, {"t": "Decimal"}, {"t": "OneParen"}],
    1: [1, {"t": "LowerAlpha"}, {"t": "OneParen"}],
    2: [1, {"t": "LowerRoman"}, {"t": "OneParen"}],
    3: [1, {"t": "Decimal"}, {"t": "TwoParens"}],
    4: [1, {"t": "LowerAlpha"}, {"t": "TwoParens"}],
    5: [1, {"t": "LowerRoman"}, {"t": "TwoParens"}],
}


def pandoc_executable() -> str:
    candidates = [
        os.environ.get("PANDOC_BIN"),
        shutil.which("pandoc"),
        "/opt/homebrew/opt/pandoc/bin/pandoc",
        "/usr/local/opt/pandoc/bin/pandoc",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    raise FileNotFoundError("pandoc not found; set PANDOC_BIN or install pandoc")


@dataclass
class SourcePara:
    inlines: List[dict]
    text: str
    quote_depth: int
    index: int


@dataclass
class Block:
    id: str
    kind: str
    level: Optional[int] = None
    inlines: List[dict] = field(default_factory=list)
    continuations: List[List[dict]] = field(default_factory=list)
    source_index: int = -1
    quote_depth: int = 0
    flags: List[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return inlines_to_text(self.inlines).strip()


@dataclass
class HeaderMetadata:
    left_text: Optional[str] = None
    code: Optional[str] = None


SEQUENCE_TYPES = (
    "decimal",
    "upper_alpha",
    "lower_alpha",
    "upper_roman",
    "lower_roman",
)


@dataclass
class LevelSpec:
    pattern: str
    sequence: str
    compiled_leading: "re.Pattern" = field(default=None, repr=False)  # type: ignore
    compiled_embedded: "re.Pattern" = field(default=None, repr=False)  # type: ignore


@dataclass
class OutlineNormalization:
    section_text: str
    source_levels: List[LevelSpec] = field(default_factory=list)


@dataclass
class DocumentParts:
    ignored_body_texts: List[str] = field(default_factory=list)
    title_texts: List[str] = field(default_factory=list)
    section_heading_texts: List[str] = field(default_factory=list)
    subheading_texts: List[str] = field(default_factory=list)
    outline_normalizations: List[OutlineNormalization] = field(default_factory=list)
    header_title_text: Optional[str] = None
    policy_code: Optional[str] = None


def inlines_to_text(inlines: List[dict]) -> str:
    out: List[str] = []
    for item in inlines:
        kind = item["t"]
        if kind == "Str":
            out.append(item["c"])
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            out.append(" ")
        elif kind in (
            "Strong",
            "Emph",
            "Underline",
            "SmallCaps",
            "Strikeout",
            "Superscript",
            "Subscript",
        ):
            out.append(inlines_to_text(item["c"]))
        elif kind == "Span":
            out.append(inlines_to_text(item["c"][1]))
        elif kind == "Quoted":
            out.append(inlines_to_text(item["c"][1]))
    return "".join(out)


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def strip_chars_from_inlines(inlines: List[dict], n: int) -> List[dict]:
    if n <= 0:
        return list(inlines)
    remaining = n
    out: List[dict] = []
    for item in inlines:
        if remaining <= 0:
            out.append(item)
            continue
        kind = item["t"]
        if kind == "Str":
            text = item["c"]
            if len(text) <= remaining:
                remaining -= len(text)
                continue
            out.append({"t": "Str", "c": text[remaining:]})
            remaining = 0
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            remaining -= 1
        elif kind in ("Strong", "Emph", "Underline", "SmallCaps", "Strikeout"):
            inner = strip_chars_from_inlines(item["c"], remaining)
            consumed = min(remaining, len(inlines_to_text(item["c"])))
            remaining -= consumed
            if inner:
                out.append({"t": kind, "c": inner})
        elif kind == "Span":
            inner = strip_chars_from_inlines(item["c"][1], remaining)
            consumed = min(remaining, len(inlines_to_text(item["c"][1])))
            remaining -= consumed
            if inner:
                out.append({"t": "Span", "c": [item["c"][0], inner]})
        else:
            out.append(item)
    while out and out[0]["t"] in ("Space", "SoftBreak"):
        out.pop(0)
    if out and out[0]["t"] == "Str":
        stripped = out[0]["c"].lstrip()
        if stripped:
            out[0] = {"t": "Str", "c": stripped}
        else:
            out.pop(0)
    return out


def split_inlines_at(inlines: List[dict], offset: int) -> Tuple[List[dict], List[dict]]:
    if offset <= 0:
        return [], list(inlines)
    left: List[dict] = []
    right: List[dict] = []
    remaining = offset
    for item in inlines:
        if remaining <= 0:
            right.append(item)
            continue
        kind = item["t"]
        if kind == "Str":
            text = item["c"]
            if len(text) <= remaining:
                left.append(item)
                remaining -= len(text)
            else:
                left.append({"t": "Str", "c": text[:remaining]})
                right.append({"t": "Str", "c": text[remaining:]})
                remaining = 0
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            left.append(item)
            remaining -= 1
        elif kind in (
            "Strong",
            "Emph",
            "Underline",
            "SmallCaps",
            "Strikeout",
            "Superscript",
            "Subscript",
        ):
            inner_text = inlines_to_text(item["c"])
            if len(inner_text) <= remaining:
                left.append(item)
                remaining -= len(inner_text)
            else:
                left_inner, right_inner = split_inlines_at(item["c"], remaining)
                if left_inner:
                    left.append({"t": kind, "c": left_inner})
                if right_inner:
                    right.append({"t": kind, "c": right_inner})
                remaining = 0
        elif kind == "Span":
            inner_text = inlines_to_text(item["c"][1])
            if len(inner_text) <= remaining:
                left.append(item)
                remaining -= len(inner_text)
            else:
                left_inner, right_inner = split_inlines_at(item["c"][1], remaining)
                if left_inner:
                    left.append({"t": "Span", "c": [item["c"][0], left_inner]})
                if right_inner:
                    right.append({"t": "Span", "c": [item["c"][0], right_inner]})
                remaining = 0
        else:
            left.append(item)
    while right and right[0]["t"] in ("Space", "SoftBreak"):
        right.pop(0)
    if right and right[0]["t"] == "Str":
        stripped = right[0]["c"].lstrip()
        if stripped:
            right[0] = {"t": "Str", "c": stripped}
        else:
            right.pop(0)
    return left, right


def trim_trailing_chars_from_inlines(inlines: List[dict], n: int) -> List[dict]:
    if n <= 0:
        return list(inlines)
    remaining = n
    out: List[dict] = []
    for item in reversed(inlines):
        if remaining <= 0:
            out.append(item)
            continue
        kind = item["t"]
        if kind == "Str":
            text = item["c"]
            if len(text) <= remaining:
                remaining -= len(text)
                continue
            out.append({"t": "Str", "c": text[:-remaining]})
            remaining = 0
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            remaining -= 1
        elif kind in (
            "Strong",
            "Emph",
            "Underline",
            "SmallCaps",
            "Strikeout",
            "Superscript",
            "Subscript",
        ):
            inner_text = inlines_to_text(item["c"])
            if len(inner_text) <= remaining:
                remaining -= len(inner_text)
                continue
            trimmed = trim_trailing_chars_from_inlines(item["c"], remaining)
            consumed = len(inner_text) - len(inlines_to_text(trimmed))
            remaining -= consumed
            if trimmed:
                out.append({"t": kind, "c": trimmed})
        elif kind == "Span":
            inner_text = inlines_to_text(item["c"][1])
            if len(inner_text) <= remaining:
                remaining -= len(inner_text)
                continue
            trimmed = trim_trailing_chars_from_inlines(item["c"][1], remaining)
            consumed = len(inner_text) - len(inlines_to_text(trimmed))
            remaining -= consumed
            if trimmed:
                out.append({"t": "Span", "c": [item["c"][0], trimmed]})
        else:
            out.append(item)
    out.reverse()
    while out and out[-1]["t"] in ("Space", "SoftBreak"):
        out.pop()
    if out and out[-1]["t"] == "Str":
        stripped = out[-1]["c"].rstrip()
        if stripped:
            out[-1] = {"t": "Str", "c": stripped}
        else:
            out.pop()
    return out


def flatten_paras(blocks: List[dict], depth: int = 0) -> List[SourcePara]:
    out: List[SourcePara] = []
    for block in blocks:
        kind = block["t"]
        if kind == "Para":
            text = normalize_text(inlines_to_text(block["c"]))
            if text:
                out.append(SourcePara(block["c"], text, depth, -1))
        elif kind == "BlockQuote":
            out.extend(flatten_paras(block["c"], depth + 1))
    for index, para in enumerate(out):
        para.index = index
    return out


def _parse_level_spec(raw: object, field_name: str) -> LevelSpec:
    if not isinstance(raw, dict):
        raise TypeError(
            f"{field_name}: each level entry must be an object with 'pattern' and "
            f"'sequence' keys, got {type(raw).__name__}"
        )
    pattern = raw.get("pattern")
    sequence = raw.get("sequence")
    if not isinstance(pattern, str) or not pattern:
        raise ValueError(f"{field_name}: missing or non-string 'pattern' in {raw!r}")
    if sequence not in SEQUENCE_TYPES:
        raise ValueError(
            f"{field_name}: 'sequence' must be one of {list(SEQUENCE_TYPES)}, "
            f"got {sequence!r}"
        )
    unexpected = set(raw) - {"pattern", "sequence"}
    if unexpected:
        raise ValueError(f"{field_name}: unknown keys: {sorted(unexpected)}")
    try:
        compiled_leading = re.compile(r"^\s*" + pattern)
        compiled_embedded = re.compile(r"(?<![A-Za-z0-9).(])" + pattern)
    except re.error as e:
        raise ValueError(f"{field_name}: invalid regex {pattern!r}: {e}")
    if compiled_leading.groups != 1:
        raise ValueError(
            f"{field_name}: pattern {pattern!r} must have exactly one capture "
            f"group (the enumeration token); got {compiled_leading.groups}"
        )
    return LevelSpec(
        pattern=pattern,
        sequence=sequence,
        compiled_leading=compiled_leading,
        compiled_embedded=compiled_embedded,
    )


def _parse_outline_normalization(raw: object, idx: int) -> OutlineNormalization:
    field_name = f"outline_normalizations[{idx}]"
    if not isinstance(raw, dict):
        raise TypeError(f"{field_name}: must be an object, got {type(raw).__name__}")
    section_text = raw.get("section_text")
    source_levels = raw.get("source_levels")
    if not isinstance(section_text, str) or not section_text.strip():
        raise ValueError(f"{field_name}: missing or empty 'section_text'")
    if not isinstance(source_levels, list) or not source_levels:
        raise ValueError(f"{field_name}: 'source_levels' must be a non-empty list")
    unexpected = set(raw) - {"section_text", "source_levels"}
    if unexpected:
        raise ValueError(f"{field_name}: unknown keys: {sorted(unexpected)}")
    levels = [
        _parse_level_spec(level, f"{field_name}.source_levels[{i}]")
        for i, level in enumerate(source_levels)
    ]
    return OutlineNormalization(section_text=section_text, source_levels=levels)


def load_document_parts(path: Path) -> DocumentParts:
    raw = json.loads(path.read_text())
    outline_raw = raw.get("outline_normalizations", [])
    if not isinstance(outline_raw, list):
        raise TypeError(
            f"outline_normalizations: must be a list, got {type(outline_raw).__name__}"
        )
    outline_normalizations = [
        _parse_outline_normalization(entry, i) for i, entry in enumerate(outline_raw)
    ]
    return DocumentParts(
        ignored_body_texts=raw.get("ignored_body_texts", []),
        title_texts=raw.get("title_texts", []),
        section_heading_texts=raw.get("section_heading_texts", []),
        subheading_texts=raw.get("subheading_texts", []),
        outline_normalizations=outline_normalizations,
        header_title_text=raw.get("header_title_text"),
        policy_code=raw.get("policy_code"),
    )


def load_source_ast(source_docx: Path) -> dict:
    pandoc_json = subprocess.check_output([pandoc_executable(), str(source_docx), "-t", "json"])
    return json.loads(pandoc_json)


def load_source_paras(source_docx: Path) -> List[SourcePara]:
    return flatten_paras(load_source_ast(source_docx)["blocks"])


def _normalize_match_string(s: str) -> str:
    chunks = s.split("\n")
    normalized = [normalize_text(chunk) for chunk in chunks]
    return "\n".join(normalized)


def _build_corpus(source_paras: List[SourcePara]) -> Tuple[str, List[Tuple[int, int, int]]]:
    parts: List[str] = []
    spans: List[Tuple[int, int, int]] = []
    cursor = 0
    for p in source_paras:
        text = p.text
        start = cursor
        end = cursor + len(text)
        spans.append((p.index, start, end))
        parts.append(text)
        cursor = end + 1
    return "\n".join(parts), spans


def _find_all(haystack: str, needle: str) -> List[int]:
    out: List[int] = []
    start = 0
    while True:
        i = haystack.find(needle, start)
        if i < 0:
            return out
        out.append(i)
        start = i + 1


def _paragraph_for_position(
    pos: int, spans: List[Tuple[int, int, int]]
) -> int:
    for idx, start, end in spans:
        if start <= pos < end:
            return idx
        if pos < start:
            return idx
    return spans[-1][0]


def _resolve_part_entry(
    entry: object,
    field_name: str,
    source_paras: List[SourcePara],
    corpus: str,
    spans: List[Tuple[int, int, int]],
) -> int:
    if not isinstance(entry, str):
        raise TypeError(
            f"{field_name}: entries must be strings (verbatim text drawn from the "
            f"target paragraph, optionally extended with neighbor paragraph text "
            f"separated by newlines), got {type(entry).__name__}"
        )
    needle = _normalize_match_string(entry)
    if not needle.strip():
        raise ValueError(f"{field_name}: empty match string is not allowed")

    positions = _find_all(corpus, needle)
    if not positions:
        raise ValueError(
            f"{field_name}: no match for {entry!r}. Provide a substring drawn "
            f"verbatim from the target paragraph (whitespace is collapsed). To "
            f"disambiguate identical paragraphs, extend the string across "
            f"paragraph boundaries with '\\n' and include some neighbor text."
        )
    if len(positions) > 1:
        targets = [_paragraph_for_position(p, spans) for p in positions]
        previews = "; ".join(
            f"#{idx}: {source_paras[idx].text[:80]}"
            f"{'...' if len(source_paras[idx].text) > 80 else ''}"
            for idx in targets
        )
        raise ValueError(
            f"{field_name}: match string {entry!r} matches {len(positions)} "
            f"positions in the corpus and is ambiguous. Extend the string with "
            f"neighbor paragraph text (using '\\n') to make it unique. "
            f"Candidate paragraphs: {previews}"
        )
    return _paragraph_for_position(positions[0], spans)


def resolve_part_strings(parts: DocumentParts, source_paras: List[SourcePara]) -> None:
    corpus, spans = _build_corpus(source_paras)
    parts.ignored_body_texts = [
        _resolve_part_entry(x, "ignored_body_texts", source_paras, corpus, spans)
        for x in parts.ignored_body_texts
    ]
    parts.title_texts = [
        _resolve_part_entry(x, "title_texts", source_paras, corpus, spans)
        for x in parts.title_texts
    ]
    parts.section_heading_texts = [
        _resolve_part_entry(x, "section_heading_texts", source_paras, corpus, spans)
        for x in parts.section_heading_texts
    ]
    parts.subheading_texts = [
        _resolve_part_entry(x, "subheading_texts", source_paras, corpus, spans)
        for x in parts.subheading_texts
    ]


_CANONICAL_LEVEL_FORMS = (
    ("decimal", "{}", ")"),
    ("lower_alpha", "{}", ")"),
    ("lower_roman", "{}", ")"),
    ("decimal", "({})", ""),
    ("lower_alpha", "({})", ""),
    ("lower_roman", "({})", ""),
)


def _index_to_alpha(n: int, upper: bool = False) -> str:
    if n < 1:
        raise ValueError(f"alpha index must be >= 1, got {n}")
    out = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        out = chr(ord("a") + rem) + out
    return out.upper() if upper else out


_ROMAN_PAIRS = (
    (1000, "m"), (900, "cm"), (500, "d"), (400, "cd"),
    (100, "c"), (90, "xc"), (50, "l"), (40, "xl"),
    (10, "x"), (9, "ix"), (5, "v"), (4, "iv"),
    (1, "i"),
)


def _index_to_roman(n: int, upper: bool = False) -> str:
    if n < 1:
        raise ValueError(f"roman index must be >= 1, got {n}")
    out = ""
    for value, numeral in _ROMAN_PAIRS:
        while n >= value:
            out += numeral
            n -= value
    return out.upper() if upper else out


def _format_sequence_token(sequence: str, n: int) -> str:
    if sequence == "decimal":
        return str(n)
    if sequence == "lower_alpha":
        return _index_to_alpha(n, upper=False)
    if sequence == "upper_alpha":
        return _index_to_alpha(n, upper=True)
    if sequence == "lower_roman":
        return _index_to_roman(n, upper=False)
    if sequence == "upper_roman":
        return _index_to_roman(n, upper=True)
    raise ValueError(f"unknown sequence type: {sequence!r}")


def canonical_marker_for_level(level_idx: int, counter: int) -> str:
    if level_idx < 0 or level_idx >= len(_CANONICAL_LEVEL_FORMS):
        raise ValueError(
            f"canonical marker only defined for levels 0..{len(_CANONICAL_LEVEL_FORMS) - 1}, "
            f"got level {level_idx}"
        )
    sequence, template, suffix = _CANONICAL_LEVEL_FORMS[level_idx]
    token = _format_sequence_token(sequence, counter)
    return template.format(token) + suffix


def _split_inlines_preserve(
    inlines: List[dict], offset: int
) -> Tuple[List[dict], List[dict]]:
    """Like split_inlines_at but does not strip leading whitespace from the
    right side. Used by Rule 0 marker rewriting which must preserve every
    character verbatim around the replacement."""
    if offset <= 0:
        return [], list(inlines)
    left: List[dict] = []
    right: List[dict] = []
    remaining = offset
    for item in inlines:
        if remaining <= 0:
            right.append(item)
            continue
        kind = item["t"]
        if kind == "Str":
            text = item["c"]
            if len(text) <= remaining:
                left.append(item)
                remaining -= len(text)
            else:
                left.append({"t": "Str", "c": text[:remaining]})
                right.append({"t": "Str", "c": text[remaining:]})
                remaining = 0
        elif kind in ("Space", "SoftBreak", "LineBreak"):
            if remaining >= 1:
                left.append(item)
                remaining -= 1
            else:
                right.append(item)
        else:
            inner_text = inlines_to_text([item])
            if len(inner_text) <= remaining:
                left.append(item)
                remaining -= len(inner_text)
            else:
                right.append(item)
                remaining = 0
    return left, right


def _replace_text_span_in_inlines(
    inlines: List[dict], offset: int, length: int, replacement: str
) -> List[dict]:
    if length <= 0:
        return inlines
    left, rest = _split_inlines_preserve(inlines, offset)
    _, right = _split_inlines_preserve(rest, length)
    out = list(left)
    if replacement:
        out.append({"t": "Str", "c": replacement})
    out.extend(right)
    return out


def _scan_marker_matches(
    text: str, source_levels: List[LevelSpec]
) -> List[Tuple[int, int, int]]:
    """Return list of (start, end, level_idx) for non-overlapping marker matches in text."""
    raw_matches: List[Tuple[int, int, int]] = []
    for level_idx, level in enumerate(source_levels):
        for m in level.compiled_embedded.finditer(text):
            raw_matches.append((m.start(), m.end(), level_idx))
    leading_match: Optional[Tuple[int, int, int]] = None
    for level_idx, level in enumerate(source_levels):
        m = level.compiled_leading.match(text)
        if m and (leading_match is None or level_idx < leading_match[2]):
            leading_match = (m.start(), m.end(), level_idx)
    if leading_match is not None:
        raw_matches = [
            (s, e, lv) for (s, e, lv) in raw_matches if s >= leading_match[1]
        ]
        raw_matches.append(leading_match)
    raw_matches.sort(key=lambda x: (x[0], x[2]))
    chosen: List[Tuple[int, int, int]] = []
    last_end = -1
    for start, end, lv in raw_matches:
        if start < last_end:
            continue
        chosen.append((start, end, lv))
        last_end = end
    return chosen


def _rewrite_paragraph_outline(
    para: SourcePara, source_levels: List[LevelSpec], counters: List[int]
) -> None:
    text = inlines_to_text(para.inlines)
    matches = _scan_marker_matches(text, source_levels)
    if not matches:
        return
    inlines = para.inlines
    delta = 0
    for start, end, level_idx in matches:
        counters[level_idx] += 1
        for j in range(level_idx + 1, len(counters)):
            counters[j] = 0
        canonical = canonical_marker_for_level(level_idx, counters[level_idx]) + " "
        adj_start = start + delta
        adj_end = end + delta
        original_len = adj_end - adj_start
        is_leading = start == 0 or text[:start].strip() == ""
        if not is_leading:
            canonical_text = " " + canonical.rstrip() + " "
        else:
            canonical_text = canonical
        inlines = _replace_text_span_in_inlines(
            inlines, adj_start, original_len, canonical_text
        )
        delta += len(canonical_text) - original_len
    para.inlines[:] = inlines
    para.text = normalize_text(inlines_to_text(para.inlines))


def normalize_outlines(parts: DocumentParts, source_paras: List[SourcePara]) -> None:
    if not parts.outline_normalizations:
        return
    section_indexes = sorted(
        idx for idx in parts.section_heading_texts if isinstance(idx, int)
    )
    corpus, spans = _build_corpus(source_paras)
    for n_idx, norm in enumerate(parts.outline_normalizations):
        target_idx = _resolve_part_entry(
            norm.section_text,
            f"outline_normalizations[{n_idx}].section_text",
            source_paras,
            corpus,
            spans,
        )
        next_idx = next(
            (i for i in section_indexes if i > target_idx), len(source_paras)
        )
        counters = [0] * len(norm.source_levels)
        for para in source_paras[target_idx + 1 : next_idx]:
            _rewrite_paragraph_outline(para, norm.source_levels, counters)


def resolve_document_parts(
    parts_in: Optional[Path] = None,
) -> DocumentParts:
    if parts_in is None:
        raise ValueError(
            "no document-parts manifest was provided. "
            "This formatter is intended to be driven by Claude Code: "
            "inspect the source document, identify the title, section headings, "
            "subheadings, running-header text, and policy code, then rerun with --parts-in /path/to/parts.json."
        )
    return load_document_parts(parts_in)


def normalize_outlines_to_docx(
    source_docx: Path,
    output_docx: Path,
    *,
    parts_in: Optional[Path] = None,
) -> None:
    """Standalone Rule 0: rewrite non-canonical outline markers in the source
    document and write the result back as a new DOCX via pandoc round-trip.
    Intended for use as a pre-pass before rule_1.py."""
    parts = resolve_document_parts(parts_in=parts_in)
    ast = load_source_ast(source_docx)
    source_paras = flatten_paras(ast["blocks"])
    resolve_part_strings(parts, source_paras)
    normalize_outlines(parts, source_paras)
    json_payload = json.dumps(ast)
    subprocess.run(
        [pandoc_executable(), "-f", "json", "-t", "docx", "-o", str(output_docx)],
        input=json_payload,
        text=True,
        check=True,
    )


_PAREN_DECIMAL_SCAN = re.compile(r"(?<![A-Za-z0-9)])\((\d+)\)(?:\s+|(?=[A-Z(])|$)")


def collect_valid_paren_decimals(source_paras: List[SourcePara]) -> set:
    universe: set = set()
    for p in source_paras:
        for m in _PAREN_DECIMAL_SCAN.finditer(p.text):
            try:
                universe.add(int(m.group(1)))
            except ValueError:
                pass
    return {n for n in universe if (n - 1) in universe or (n + 1) in universe}


def parse_leading_marker(
    inlines: List[dict],
    prev_l1: Optional[str],
    prev_l4: Optional[str],
    valid_paren_decimals: Optional[set] = None,
) -> Tuple[Optional[int], Optional[str], List[dict], bool]:
    text = inlines_to_text(inlines)
    matches = []
    for level, regex in enumerate(LEVEL_PATTERNS):
        match = regex.match(text)
        if match:
            if level == 3 and valid_paren_decimals is not None:
                try:
                    if int(match.group(1)) not in valid_paren_decimals:
                        continue
                except ValueError:
                    continue
            matches.append((level, match.end(), match.group(1)))
    if not matches:
        return None, None, inlines, False
    if len(matches) == 1:
        level, end, token = matches[0]
        return level, token, strip_chars_from_inlines(inlines, end), False

    roman_matches = [match for match in matches if match[0] in (2, 5)]
    alpha_matches = [match for match in matches if match[0] in (1, 4)]

    if roman_matches and any(len(match[2]) > 1 for match in roman_matches):
        level, end, token = next(match for match in roman_matches if len(match[2]) > 1)
        return level, token, strip_chars_from_inlines(inlines, end), False

    token = matches[0][2]
    if token == "i":
        if prev_l1 == "h":
            level, end, token = next(match for match in alpha_matches if match[0] == 1)
            return level, token, strip_chars_from_inlines(inlines, end), False
        if prev_l4 == "h":
            level, end, token = next(match for match in alpha_matches if match[0] == 4)
            return level, token, strip_chars_from_inlines(inlines, end), False
    if token == "v":
        if prev_l1 == "u":
            level, end, token = next(match for match in alpha_matches if match[0] == 1)
            return level, token, strip_chars_from_inlines(inlines, end), False
        if prev_l4 == "u":
            level, end, token = next(match for match in alpha_matches if match[0] == 4)
            return level, token, strip_chars_from_inlines(inlines, end), False
    if token == "x":
        if prev_l1 == "w":
            level, end, token = next(match for match in alpha_matches if match[0] == 1)
            return level, token, strip_chars_from_inlines(inlines, end), False
        if prev_l4 == "w":
            level, end, token = next(match for match in alpha_matches if match[0] == 4)
            return level, token, strip_chars_from_inlines(inlines, end), False

    ambiguous = len(alpha_matches) > 0 and len(roman_matches) > 0
    level, end, token = roman_matches[0] if roman_matches else matches[0]
    return level, token, strip_chars_from_inlines(inlines, end), ambiguous


def find_embedded_marker_offset(
    inlines: List[dict],
    valid_paren_decimals: Optional[set] = None,
) -> Optional[int]:
    text = inlines_to_text(inlines)
    best: Optional[int] = None
    for regex in EMBEDDED_MARKER_PATTERNS:
        for match in regex.finditer(text):
            if match.start() <= 0:
                continue
            token = match.group(1) if match.lastindex else match.group(0)
            if valid_paren_decimals is not None:
                paren_digit = re.match(r"\((\d+)\)", token)
                if paren_digit and int(paren_digit.group(1)) not in valid_paren_decimals:
                    continue
            if best is None or match.start() < best:
                best = match.start()
    return best


def cleanup_left_split_fragment(inlines: List[dict]) -> List[dict]:
    text = inlines_to_text(inlines)
    new_text = re.sub(r"\s*\(\s*$", "", text)
    new_text = re.sub(r"\s+(and|or)\s*$", "", new_text, flags=re.IGNORECASE)
    new_text = re.sub(r"\s+$", "", new_text)
    trim = len(text) - len(new_text)
    if trim <= 0:
        return inlines
    return trim_trailing_chars_from_inlines(inlines, trim)


def split_on_embedded_markers(
    source_paras: List[SourcePara],
    valid_paren_decimals: Optional[set] = None,
) -> List[SourcePara]:
    out: List[SourcePara] = []
    for para in source_paras:
        pending = [para.inlines]
        while pending:
            current_inlines = pending.pop(0)
            offset = find_embedded_marker_offset(current_inlines, valid_paren_decimals)
            if offset is None:
                out.append(
                    SourcePara(
                        current_inlines,
                        normalize_text(inlines_to_text(current_inlines)),
                        para.quote_depth,
                        para.index,
                    )
                )
                continue
            left, right = split_inlines_at(current_inlines, offset)
            left = cleanup_left_split_fragment(left)
            left_text = normalize_text(inlines_to_text(left))
            right_text = normalize_text(inlines_to_text(right))
            if not left_text or not right_text:
                out.append(
                    SourcePara(
                        current_inlines,
                        normalize_text(inlines_to_text(current_inlines)),
                        para.quote_depth,
                        para.index,
                    )
                )
                continue
            pending.insert(0, right)
            pending.insert(0, left)
    return out


def merge_or_continue(prev_block: Block, para: SourcePara) -> None:
    prev_text = prev_block.text
    starts_lower = para.text[:1].islower()
    ends_mid_sentence = bool(prev_text) and not re.search(r"[.!?;:]$", prev_text)
    if starts_lower or ends_mid_sentence:
        prev_block.inlines.append({"t": "Space"})
        prev_block.inlines.extend(para.inlines)
        return
    prev_block.continuations.append(para.inlines)


def classify(
    source_paras: List[SourcePara],
    parts: DocumentParts,
    valid_paren_decimals: Optional[set] = None,
) -> List[Block]:
    blocks: List[Block] = []
    ignored_indexes = set(parts.ignored_body_texts)
    title_indexes = set(parts.title_texts)
    section_indexes = set(parts.section_heading_texts)
    subheading_indexes = set(parts.subheading_texts)
    prev_l1: Optional[str] = None
    prev_l4: Optional[str] = None
    last_list: Optional[Block] = None
    next_id = 0

    for para in source_paras:
        text = para.text.strip()
        if not text or para.index in ignored_indexes:
            continue

        if para.index in title_indexes:
            blocks.append(
                Block(
                    f"b{next_id:04d}",
                    "title",
                    inlines=para.inlines,
                    source_index=para.index,
                    quote_depth=para.quote_depth,
                )
            )
            next_id += 1
            last_list = None
            prev_l1 = None
            prev_l4 = None
            continue

        if para.index in section_indexes:
            blocks.append(
                Block(
                    f"b{next_id:04d}",
                    "section_heading",
                    inlines=para.inlines,
                    source_index=para.index,
                    quote_depth=para.quote_depth,
                )
            )
            next_id += 1
            last_list = None
            prev_l1 = None
            prev_l4 = None
            continue

        level, token, stripped, ambiguous = parse_leading_marker(
            para.inlines, prev_l1, prev_l4, valid_paren_decimals
        )
        if level is not None:
            blocks.append(
                Block(
                    id=f"b{next_id:04d}",
                    kind="list_item",
                    level=level,
                    inlines=stripped,
                    source_index=para.index,
                    quote_depth=para.quote_depth,
                    flags=["ambiguous_marker"] if ambiguous else [],
                )
            )
            next_id += 1
            last_list = blocks[-1]
            if level == 1:
                prev_l1 = token
            elif level == 4:
                prev_l4 = token
            continue

        if para.index in subheading_indexes:
            blocks.append(
                Block(
                    f"b{next_id:04d}",
                    "subheading",
                    inlines=para.inlines,
                    source_index=para.index,
                    quote_depth=para.quote_depth,
                )
            )
            next_id += 1
            last_list = None
            prev_l1 = None
            prev_l4 = None
            continue

        if last_list is not None and (para.quote_depth > 0 or last_list.level is not None):
            merge_or_continue(last_list, para)
            continue

        blocks.append(
            Block(
                f"b{next_id:04d}",
                "body",
                inlines=para.inlines,
                source_index=para.index,
                quote_depth=para.quote_depth,
            )
        )
        next_id += 1
        last_list = None
        prev_l1 = None
        prev_l4 = None
    return blocks


def item_to_pandoc_blocks(block: Block) -> List[dict]:
    if not block.continuations:
        return [{"t": "Para", "c": block.inlines}]
    merged = list(block.inlines)
    for continuation in block.continuations:
        merged.extend([{"t": "LineBreak"}, {"t": "LineBreak"}])
        merged.extend(continuation)
    return [{"t": "Para", "c": merged}]


def compose_list(blocks: List[Block], index: int) -> Tuple[dict, int]:
    base_level = blocks[index].level or 0
    items: List[List[dict]] = []
    while index < len(blocks) and blocks[index].kind == "list_item" and (blocks[index].level or 0) >= base_level:
        current = blocks[index]
        if (current.level or 0) == base_level:
            item_blocks = item_to_pandoc_blocks(current)
            items.append(item_blocks)
            index += 1
            if index < len(blocks) and blocks[index].kind == "list_item" and (blocks[index].level or 0) > base_level:
                nested, index = compose_list(blocks, index)
                items[-1].append(nested)
        else:
            nested, index = compose_list(blocks, index)
            if items:
                items[-1].append(nested)
            else:
                items.append([nested])
    return {"t": "OrderedList", "c": [LEVEL_TO_OL_ATTR[base_level], items]}, index


def compose_blocks_ast(blocks: List[Block]) -> dict:
    pandoc_blocks: List[dict] = []
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.kind == "title":
            pandoc_blocks.append({"t": "Header", "c": [1, ["", [], []], block.inlines]})
            index += 1
        elif block.kind == "section_heading":
            pandoc_blocks.append({"t": "Header", "c": [2, ["", [], []], block.inlines]})
            index += 1
        elif block.kind == "subheading":
            pandoc_blocks.append({"t": "Header", "c": [3, ["", [], []], block.inlines]})
            index += 1
        elif block.kind == "list_item":
            ordered_list, index = compose_list(blocks, index)
            pandoc_blocks.append(ordered_list)
        else:
            pandoc_blocks.append({"t": "Para", "c": block.inlines})
            index += 1
    return {"pandoc-api-version": [1, 23, 1], "meta": {}, "blocks": pandoc_blocks}


def render(ast_path: Path, output_docx: Path) -> None:
    subprocess.check_call([pandoc_executable(), str(ast_path), "-f", "json", "-t", "docx", "-o", str(output_docx)])


def render_blocks_to_docx(
    blocks: List[Block],
    output_docx: Path,
    *,
    ast_out: Optional[Path] = None,
) -> dict:
    ast = compose_blocks_ast(blocks)
    ast_path = ast_out or output_docx.with_suffix(".ast.json")
    ast_path.write_text(json.dumps(ast) + "\n")
    render(ast_path, output_docx)
    return ast


def build_text_hierarchy_docx(
    source_docx: Path,
    output_docx: Path,
    *,
    parts_in: Optional[Path] = None,
    ast_out: Optional[Path] = None,
) -> List[Block]:
    parts = resolve_document_parts(parts_in=parts_in)
    source_paras = load_source_paras(source_docx)
    resolve_part_strings(parts, source_paras)
    normalize_outlines(parts, source_paras)
    valid_paren_decimals = collect_valid_paren_decimals(source_paras)
    blocks = classify(
        split_on_embedded_markers(source_paras, valid_paren_decimals),
        parts,
        valid_paren_decimals,
    )
    render_blocks_to_docx(blocks, output_docx, ast_out=ast_out)
    apply_heading_styles_to_docx(output_docx)
    apply_body_styles_to_docx(output_docx)
    return blocks


def header_metadata_from_parts(parts: DocumentParts) -> HeaderMetadata:
    return HeaderMetadata(left_text=parts.header_title_text, code=parts.policy_code)


def extract_header_metadata_from_docx(
    *,
    parts_in: Optional[Path] = None,
) -> HeaderMetadata:
    parts = resolve_document_parts(parts_in=parts_in)
    return header_metadata_from_parts(parts)


def ensure_style(doc: Document, name: str) -> None:
    if name not in doc.styles:
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def ensure_doc_styles(doc: Document) -> None:
    ensure_style(doc, "CorgiTitle")
    ensure_style(doc, "CorgiSection")
    ensure_style(doc, "CorgiSubheading")
    ensure_style(doc, "CorgiBody")


def apply_header(doc: Document, header_meta: HeaderMetadata) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.35)
        header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        header_para.clear()
        header_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
        if header_meta.left_text and header_meta.code:
            header_text = f"{header_meta.left_text}\t{header_meta.code}"
        else:
            header_text = header_meta.left_text or header_meta.code or ""
        run = header_para.add_run(header_text)
        run.font.name = "Inter"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_header_to_docx(path: Path, header: HeaderMetadata) -> None:
    doc = Document(path)
    apply_header(doc, header)
    doc.save(path)


def apply_heading_styles(doc: Document) -> None:
    ensure_doc_styles(doc)
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Heading 1":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(18)
            for run in para.runs:
                run.font.name = "Bricolage Grotesque"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bricolage Grotesque")
                run.font.size = Pt(26)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name == "Heading 2":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(8)
            for run in para.runs:
                run.font.name = "Bricolage Grotesque"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bricolage Grotesque")
                run.font.size = Pt(14)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        elif style_name == "Heading 3":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = "Bricolage Grotesque"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Bricolage Grotesque")
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)


def apply_heading_styles_to_docx(path: Path) -> None:
    doc = Document(path)
    apply_heading_styles(doc)
    doc.save(path)


def apply_body_styles(doc: Document) -> None:
    ensure_doc_styles(doc)
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name not in ("Heading 1", "Heading 2", "Heading 3"):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.name = "Inter"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Inter")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 0)


def apply_body_styles_to_docx(path: Path) -> None:
    doc = Document(path)
    apply_body_styles(doc)
    doc.save(path)


def patch_list_marker_style(path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path, "r") as zin:
        items = [(item, zin.read(item.filename)) for item in zin.infolist()]
        root = etree.fromstring(zin.read("word/numbering.xml"))

    for level in root.xpath(".//w:lvl", namespaces=ns):
        for rpr in level.xpath("./w:rPr", namespaces=ns):
            level.remove(rpr)
        rpr = etree.SubElement(level, f"{{{ns['w']}}}rPr")
        etree.SubElement(
            rpr,
            f"{{{ns['w']}}}rFonts",
            attrib={
                f"{{{ns['w']}}}ascii": "Inter",
                f"{{{ns['w']}}}hAnsi": "Inter",
                f"{{{ns['w']}}}cs": "Inter",
                f"{{{ns['w']}}}eastAsia": "Inter",
            },
        )
        etree.SubElement(rpr, f"{{{ns['w']}}}color", attrib={f"{{{ns['w']}}}val": "000000"})
        etree.SubElement(rpr, f"{{{ns['w']}}}sz", attrib={f"{{{ns['w']}}}val": "22"})
        etree.SubElement(rpr, f"{{{ns['w']}}}szCs", attrib={f"{{{ns['w']}}}val": "22"})

    numbering = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item, data in items:
            if item.filename == "word/numbering.xml":
                data = numbering
            zout.writestr(item, data)
    tmp.replace(path)


def normalize_list_formatting(path: Path) -> None:
    patch_list_levels(path)
    set_list_suffix(path, "space")
    strip_list_ind_overrides(path)
    patch_list_marker_style(path)


def apply_page_layout_and_header(
    source_docx: Path,
    target_docx: Path,
    *,
    parts_in: Optional[Path] = None,
) -> None:
    _ = source_docx
    header = extract_header_metadata_from_docx(parts_in=parts_in)
    apply_header_to_docx(target_docx, header)
